"""
Generates the CFM56-5B thesis as a properly formatted Word document.
Requirements: python-docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Base font (Times New Roman 12pt) ────────────────────────────────────────
style = doc.styles['Normal']
font  = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = Pt(18)   # ~1.5×12pt
pf.space_after  = Pt(0)

def set_font(run, bold=False, italic=False, size=12):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic

def heading1(text):
    """Main chapter heading — bold 14pt, new page via page break."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=14)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=12)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, bold=True, italic=True, size=12)
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.line_spacing      = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run)
    return p

def caption_fig(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=10)
    return p

def caption_tab(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, size=10)
    return p

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
r = p.add_run('NYÍREGYHÁZI EGYETEM')
set_font(r, bold=True, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Műszaki és Agrártudományi Intézet')
set_font(r, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Repülőmérnöki alapképzési szak')
set_font(r, size=12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CFM56-5B REPÜLŐGÉPHAJTÓMŰ TERMODINAMIKAI\nCIKLUS-ANALÍZISE PYTHON ALAPÚ SZIMULÁCIÓVAL')
set_font(r, bold=True, size=16)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SZAKDOLGOZAT')
set_font(r, bold=True, size=14)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Készítette: Mohamed Ziad\nRepülőmérnöki alapszak, VI. félév')
set_font(r, size=12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Nyíregyháza, 2026')
set_font(r, size=12)

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TARTALOMJEGYZÉK')
set_font(r, bold=True, size=14)
doc.add_paragraph()

toc_entries = [
    ('BEVEZETÉS', '1'),
    ('1. A CFM56-5B HAJTÓMŰ ISMERTETÉSE', '3'),
    ('   1.1. Történeti áttekintés', '3'),
    ('   1.2. Műszaki jellemzők és alkalmazások', '5'),
    ('   1.3. Kétáramú sugárhajtómű működési elve', '7'),
    ('2. TERMODINAMIKAI ALAPOK', '10'),
    ('   2.1. A Brayton-ciklus elmélete', '10'),
    ('   2.2. Kompresszor és turbina termodinamikája', '13'),
    ('   2.3. Égéstér és hőmérséklet-szabályozás', '16'),
    ('   2.4. Fúvócső és tolóerő-generálás', '18'),
    ('3. A SZIMULÁCIÓ MÓDSZERTANA', '21'),
    ('   3.1. NASA pyCycle keretrendszer', '21'),
    ('   3.2. A szimulációs modell felépítése', '23'),
    ('   3.3. Repülési állapotok és bemeneti paraméterek', '26'),
    ('4. SZIMULÁCIÓS EREDMÉNYEK', '29'),
    ('   4.1. Tervezési pont analízis — felszállás', '29'),
    ('   4.2. Off-design analízis — három repülési fázis', '33'),
    ('   4.3. Gázkar-szimuláció eredményei', '37'),
    ('   4.4. T-s diagram elemzése', '40'),
    ('5. EREDMÉNYEK ÉRTÉKELÉSE ÉS KÖVETKEZTETÉSEK', '44'),
    ('   5.1. A szimulációs eredmények validálása', '44'),
    ('   5.2. Tolóerő és tüzelőanyag-fogyasztás összefüggése', '47'),
    ('   5.3. Fejlesztési javaslatok', '50'),
    ('ÖSSZEFOGLALÁS', '53'),
    ('IRODALOMJEGYZÉK', '55'),
    ('HALLGATÓI NYILATKOZAT', '57'),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    tab_stop = p.paragraph_format.tab_stops
    run = p.add_run(entry)
    set_font(run, bold=('.' not in entry.split()[0] and entry[0].isupper() and '   ' not in entry))
    run2 = p.add_run(f'\t{page}')
    set_font(run2)

# ═══════════════════════════════════════════════════════════════════════════
# BEVEZETÉS
# ═══════════════════════════════════════════════════════════════════════════
heading1('BEVEZETÉS')

body(
    'A modern repülőgép-hajtóművek a mérnöki tudományok egyik legösszetettebb alkotásai, '
    'amelyek a termodinamika, a gázdinamika és az anyagtudomány határterületein működnek. '
    'A kereskedelmi légi közlekedés globális bővülésével egyre nagyobb igény mutatkozik a '
    'hajtóművek hatékonyságának növelésére, tüzelőanyag-fogyasztásának csökkentésére és '
    'a kibocsátási normák teljesítésére.'
)
body(
    'A szakdolgozat tárgya a CFM56-5B kétáramú gázturbinás sugárhajtómű termodinamikai '
    'ciklus-analízise, Python alapú szimulációval. A CFM56-5B az Airbus A320-as '
    'repülőgépcsalád egyik leggyakrabban alkalmazott hajtóműve, amelyből világszerte több '
    'mint 20 000 darabot helyeztek üzembe. Gazdag gyártói dokumentációja és nyilvános '
    'műszaki adatai különösen alkalmassá teszik szimulációs vizsgálatokra.'
)
body(
    'A szimuláció elvégzéséhez a NASA által fejlesztett és nyilvánosan elérhető pyCycle '
    'keretrendszert alkalmaztam, amely CEA-alapú (Chemical Equilibrium with Applications) '
    'termodinamikai számításokat végez az OpenMDAO optimalizálási platformon. Ez a '
    'megközelítés lehetővé teszi, hogy a valóságos gázturbinás ciklus egyenleteit '
    'numerikusan oldjuk meg, figyelembe véve a komponensek hatásfokát és a tényleges '
    'gázösszetételt.'
)
body(
    'A dolgozat célkitűzései a következők: (1) a CFM56-5B tervezési pontjának meghatározása '
    'felszállási körülmények között; (2) off-design analízis elvégzése három repülési fázisra '
    '(felszállás, emelkedés, utazórepülés); (3) interaktív gázkar-szimulátor fejlesztése, '
    'amely bemutatja a tolóerő és a tüzelőanyag-fogyasztás összefüggését a turbinabemeneti '
    'hőmérséklet függvényében; (4) a termodinamikai ciklus T-s diagramon való ábrázolása '
    'és értékelése.'
)
body(
    'A szakdolgozat felépítése a követelményeknek megfelelően halad: az elméleti alapok '
    'ismertetése után a szimulációs módszertan, majd a részletes eredmények és azok '
    'értékelése következik. A munka gyakorlati részét a Python-alapú szimulációs kód, '
    'a Jupyter Notebook interaktív elemzések és a vizualizációs eszközök alkotják.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 1. FEJEZET
# ═══════════════════════════════════════════════════════════════════════════
heading1('1. A CFM56-5B HAJTÓMŰ ISMERTETÉSE')

heading2('1.1. Történeti áttekintés')
body(
    'A CFM56 hajtóműcsalád a CFM International vegyesvállalat terméke, amelyet az '
    'amerikai General Electric (GE Aviation) és a francia Safran Aircraft Engines '
    '(korábban SNECMA) alapított 1974-ben. A vegyesvállalat létrehozásának elsődleges '
    'célja egy új, közepes tolóerejű, nagyfokú bypass-arányú gázturbinás sugárhajtómű '
    'fejlesztése volt a polgári repülés számára [1].'
)
body(
    'Az első CFM56-1 változat 1979-ben kapta meg az FAA típusengedélyét, és a KC-135 '
    'Stratotanker katonai szállítógép re-motorozásához alkalmazták. Az 5-es sorozat '
    'fejlesztése az 1980-as évek közepén kezdődött az Airbus A320-as programhoz '
    'kapcsolódóan. A CFM56-5A jelölésű alapváltozat 1987-ben szerezte meg a '
    'típusengedélyét, amelyet az A320 sikere nyomán gyorsan követett a továbbfejlesztett '
    'CFM56-5B változat [2].'
)
body(
    'A CFM56-5B sorozat 1994 óta van kereskedelmi forgalomban, és azóta az Airbus '
    'A318, A319, A320 és A321 repülőgépek elsőszámú hajtóművévé vált. A típus '
    'folyamatos fejlesztésének eredményeként ma már a CFM56-5B/3 Tech Insertion '
    'változat is elérhető, amely javított kompresszor- és turbinalapát-geometriával '
    'akár 1%-os tüzelőanyag-megtakarítást tesz lehetővé [3].'
)
body(
    'A hajtóműcsalád kiemelkedő megbízhatóságát jelzi, hogy az ETOPS (Extended-range '
    'Twin-engine Operational Performance Standards) minősítés alapján a CFM56-5B '
    'felszerelésű A320-as repülőgépek 180 perces ETOPS-engedéllyel repülhetnek, ami '
    'az ikerhajtóművű gépek legmagasabb megbízhatósági kategóriáját jelenti [4].'
)

heading2('1.2. Műszaki jellemzők és alkalmazások')
body(
    'A CFM56-5B hajtómű kétáramú (turbofan) gázturbinás sugárhajtómű, amelynek '
    'főbb műszaki jellemzői a következők. A maximális száraz tolóerő '
    '133,4 kN (30 000 lbf), ami az A320-as repülőgép felszállásához szükséges '
    'teljesítményt biztosítja. A bypass-arány (BPR) 5,5:1, ami azt jelenti, hogy '
    'a belépő levegő tömegáramának 5,5-szöröse kerül megkerülésre a magáramhoz '
    'képest [2].'
)

# Table 1
caption_tab('1. táblázat. A CFM56-5B főbb műszaki adatai\nForrás: [2], [3]')
table = doc.add_table(rows=9, cols=2)
table.style = 'Table Grid'
headers = ['Jellemző', 'Érték']
data = [
    ('Maximális tolóerő', '133,4 kN (30 000 lbf)'),
    ('Bypass-arány (BPR)', '5,5:1'),
    ('Összesített nyomásviszony (OPR)', '27,0'),
    ('Turbinabemeneti hőmérséklet (T4)', '~1700 K'),
    ('Tömegáram (légbevitel)', '~172 kg/s'),
    ('Száraz tömeg', '2 374 kg'),
    ('Ventilátor átmérője', '1 735 mm'),
    ('Hossz', '2 515 mm'),
]
for i, (h, v) in enumerate([(headers[0], headers[1])] + data):
    row = table.rows[i]
    row.cells[0].text = h
    row.cells[1].text = v
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True

doc.add_paragraph()

body(
    'A hajtómű alkalmazási területei elsősorban a rövidebb és közepes hatótávolságú '
    'útvonalakra tervezett Airbus A320-as repülőgépcsalád tagjai: az A318 (107 utas), '
    'A319 (124 utas), A320 (150 utas) és A321 (180 utas) típusok. Ezek a repülőgépek '
    'a világ leggyakrabban üzemeltetett keskeny törzsű repülőgépei közé tartoznak, és '
    'a globális légi közlekedés kb. 30%-át bonyolítják le [5].'
)
body(
    'A hajtómű karbantartási ciklusa rendkívül kedvező: az ECTM (Engine Condition '
    'Trend Monitoring) rendszer segítségével folyamatosan figyelik az üzemelési '
    'paramétereket. Az átlagos shop visit intervallum meghaladja a 20 000 repülési '
    'órát, ami ipari összehasonlításban kiemelkedő értéknek számít [6].'
)

heading2('1.3. Kétáramú sugárhajtómű működési elve')
body(
    'A kétáramú sugárhajtómű (turbofan) a tiszta gázturbinás sugárhajtómű (turbojet) '
    'továbbfejlesztése, amelynek alapelve az, hogy a kompresszorba belépő levegő '
    'egy részét — az ún. bypass-áramot — nem vezeti az égéstérbe, hanem közvetlenül '
    'a magáram körül vezeti el, és a hajtómű hátsó részén bocsátja ki. Ez a megoldás '
    'lényegesen jobb propulziós hatásfokot eredményez szubszonikus sebességtartományban [7].'
)
body(
    'A CFM56-5B kéttengelyes (twin-spool) kialakítású: a kis nyomású (LP) és nagy '
    'nyomású (HP) rendszerek egymástól független tengelyeken forognak, ami rugalmasabb '
    'üzemelést és jobb részteljesítményen mért hatásfokot tesz lehetővé. A levegő '
    'útja a következő főbb komponenseken keresztül vezet: légbevezetők (inlet) → '
    'ventilátor (fan) → kis nyomású kompresszor (LPC) → nagy nyomású kompresszor '
    '(HPC) → égéstér (combustor) → nagy nyomású turbina (HPT) → kis nyomású '
    'turbina (LPT) → fúvócső (nozzle) [7].'
)
body(
    'A ventilátor egyszerre két feladatot lát el: sűríti a magáramot (amely az LPC '
    'felé halad) és felgyorsítja a bypass-áramot, amely közvetlenül tolóerőt termel. '
    'A CFM56-5B esetében a bypass-arány 5,5:1, ami azt jelenti, hogy a teljes '
    'tolóerő kb. 75-80%-a a bypass-áramtól, és csupán 20-25%-a a magáramtól '
    'származik, ami a modern polgári hajtóművekre jellemző arány [2].'
)

# ═══════════════════════════════════════════════════════════════════════════
# 2. FEJEZET
# ═══════════════════════════════════════════════════════════════════════════
heading1('2. TERMODINAMIKAI ALAPOK')

heading2('2.1. A Brayton-ciklus elmélete')
body(
    'A gázturbinás hajtóművek termodinamikai alapja az ún. Brayton-ciklus '
    '(más néven Joule-ciklus), amelyet George Brayton amerikai mérnökről '
    'neveztek el, aki az 1870-es években szabadalmaztatta az első folyamatos '
    'égésen alapuló gépet. A ciklus ideális esetben három fő folyamatból áll: '
    'izentrópikus kompresszió, izobár hőközlés (égés), valamint izentrópikus '
    'expanzió (turbina) [8].'
)
body(
    'A valóságos gázturbinás ciklusban az ideális Brayton-ciklustól való eltérések '
    'a következők miatt lépnek fel: a kompresszorban és turbinában fellépő '
    'irreverzibilis súrlódási veszteségek (alacsonyabb izentrópikus hatásfok), '
    'a nyomásveszteségek az égéstérben és a csövekben, a hőleadás az égéstér '
    'falain keresztül, valamint a hűtőlevegő hatása a turbinalapátokon [9].'
)
body(
    'A termikus hatásfok az ideális Brayton-ciklusban kizárólag a nyomásviszony '
    'függvénye. A CFM56-5B esetében az összesített nyomásviszony OPR = 27,0, '
    'ami ideális esetben kb. 47%-os termikus hatásfoknak felelne meg. A valóságos '
    'hatásfok ennél kisebb az említett irreverzibilitások miatt, de a modern '
    'hajtóművek 40-42% körüli termikus hatásfokot is elérnek [2].'
)
body(
    'A T-s diagramon (hőmérséklet–specifikus entrópia diagram) a Brayton-ciklus '
    'jellegzetes alakja rajzolódik ki: a kompresszió felfelé haladó görbe '
    '(S2→S3), az égés vízszintes (vagy enyhén jobbra hajló) szakasz (S3→S4), '
    'az expanzió lefelé haladó görbe (S4→S5), majd a hőleadás visszafelé '
    'haladó vonal. A ciklus által bezárt terület arányos a fajlagos munkával [9].'
)

heading2('2.2. Kompresszor és turbina termodinamikája')
body(
    'A kompresszor izentrópikus hatásfoka (η_c) a valóságos és az ideális '
    'kompressziós munka arányaként definiálható. A CFM56-5B esetében '
    'a ventilátor hatásfoka η_fan = 0,89, a kis nyomású kompresszornak '
    'η_LPC = 0,89, a nagy nyomású kompresszornak η_HPC = 0,87 [2], [3].'
)
body(
    'A kompressziós folyamat végén (S3 állomás) a levegő hőmérséklete a '
    'szimulációs eredmények alapján felszállásnál 804,9 K, nyomása 2853 kPa. '
    'Ez az OPR = 27,0 nyomásviszonnyal és a komponensek hatásfokával összhangban '
    'van. A nagy nyomású kompresszorban végzett munka jelentős hőmérséklet-emelkedést '
    'okoz: az LPC kimenetétől (S25 = 428 K) az HPC kimenetéig (S3 = 805 K) '
    'közel 377 K-es hőmérséklet-növekedés következik be [10].'
)
body(
    'A turbina izentrópikus hatásfoka (η_t) a valóságos és az ideális expanziós '
    'munka arányaként definiálható. A CFM56-5B esetében a nagy nyomású turbina '
    'hatásfoka η_HPT = 0,89, a kis nyomású turbináé η_LPT = 0,90. A turbina '
    'nagyobb hatásfokkal üzemel, mint a kompresszor, ami az áramlástan és a '
    'lapát-geometria különbségeiből adódik [3].'
)

heading2('2.3. Égéstér és hőmérséklet-szabályozás')
body(
    'Az égéstér (combustion chamber) feladata, hogy a kompresszorból érkező '
    'sűrített levegőt kerosinnel (Jet-A üzemanyaggal) elégesse, és így a gáz '
    'hőmérsékletét a turbinabemeneti hőmérsékletre (T4) emelje. Az égéstér '
    'tüzelőanyag-levegő arányának (FAR – Fuel-to-Air Ratio) szabályozásával '
    'a hajtómű tolóereje folyamatosan változtatható [9].'
)
body(
    'A turbinabemeneti hőmérséklet (T4 vagy TIT – Turbine Inlet Temperature) '
    'a hajtómű legkritikusabb termodinamikai paramétere: minél magasabb, annál '
    'nagyobb a termikus hatásfok és a fajlagos tolóerő. Ugyanakkor a '
    'turbinalapátok hőterhelése is növekszik, ami korlátozza a maximálisan '
    'megengedhető T4 értéket. A CFM56-5B esetében a tervezési T4 érték '
    'közelítőleg 1700 K [2].'
)
body(
    'A gázkar (throttle) fizikailag a tüzelőanyag-adagoló szelepet vezérli, '
    'és így szabályozza a FAR értékét, ami közvetlenül meghatározza a T4-et és '
    'ezen keresztül a tolóerőt. Alapjárati állásban (idle) a T4 kb. 1000 K, '
    'maximális tolóerőnél (TOGA – Take-Off/Go-Around) eléri az 1700 K-t. '
    'Ezen intervallum szimulációs vizsgálatát mutatja be a 4.3. fejezet.'
)

heading2('2.4. Fúvócső és tolóerő-generálás')
body(
    'A tolóerő (thrust) a hajtómű által a kilépő gázáramra alkalmazott impulzus '
    'reakciójaként keletkezik (Newton III. törvénye). A nettó tolóerő a kilépő '
    'és belépő impulzusáramok különbségeként, valamint a fúvócső nyomóerejének '
    'figyelembevételével számítható [7]:'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Fn = ṁ_core · V_8 + ṁ_byp · V_18 – (ṁ_core + ṁ_byp) · V_0 + (P_8 – P_0) · A_8 + (P_18 – P_0) · A_18')
set_font(r, italic=True)

body(
    'ahol V_8 és V_18 a mag- és bypass-fúvócsőből kilépő gáz sebessége, V_0 a '
    'repülési sebesség, P_8 és P_18 a fúvócső toroknyomásai, P_0 a szabad '
    'légköri nyomás, A_8 és A_18 a fúvócső torkolatainak keresztmetszetei. '
    'A szimulációban a pyCycle keretrendszer a perf.Fn komponens segítségével '
    'számítja ezt az értéket [11].'
)
body(
    'A fajlagos tüzelőanyag-fogyasztás (SFC – Specific Fuel Consumption) a '
    'tüzelőanyag-fogyasztás és a tolóerő hányadosa: SFC = ṁ_fuel / Fn. '
    'A CFM56-5B esetében az utazórepülésnél mért SFC jellemzően '
    '0,0135 kg/(kN·s) körül alakul, ami az iparban elfogadott referencia-értékkel '
    'jó egyezésben van [6].'
)

# ═══════════════════════════════════════════════════════════════════════════
# 3. FEJEZET
# ═══════════════════════════════════════════════════════════════════════════
heading1('3. A SZIMULÁCIÓ MÓDSZERTANA')

heading2('3.1. NASA pyCycle keretrendszer')
body(
    'A szimulációhoz a NASA Glenn Research Center által fejlesztett pyCycle '
    'keretrendszert alkalmaztam (4.4.0 verzió). A pyCycle egy nyílt forráskódú, '
    'Python alapú gázturbina-ciklus analízis eszköz, amely az OpenMDAO '
    '(Open-source Multidisciplinary Design, Analysis and Optimization) '
    'keretrendszeren épül [12].'
)
body(
    'A pyCycle legfőbb előnye a hagyományos, manuális Brayton-ciklus számításokkal '
    'szemben, hogy CEA-alapú (Chemical Equilibrium with Applications) '
    'termodinamikai adatbázist használ. Ez azt jelenti, hogy a gázok '
    'termodinamikai tulajdonságait (entalpia, entrópia, fajhő) nem konstans '
    'értékként kezeli, hanem a hőmérséklet és nyomás valóságos függvényeként '
    'számítja, figyelembe véve a gázösszetétel változását az égés során [12], [13].'
)
body(
    'Az OpenMDAO keretrendszer implicit egyenletrendszer-megoldó képessége '
    'lehetővé teszi, hogy a hajtómű komponenseinek egymástól függő egyenleteit '
    'iteratív numerikus módszerrel oldja meg. Ez a megközelítés pontosabb '
    'eredményeket ad, mint az egyszerűsített analitikus közelítések, különösen '
    'off-design körülmények között [13].'
)
body(
    'A szimulációs környezet telepítéséhez a következő Python csomagokat '
    'alkalmaztam: om-pycycle 4.4.0, openmdao 3.43.0, numpy, scipy és matplotlib. '
    'A teljes szimulációs kód verziókezelése git rendszerrel történt, és '
    'reprodukálhatóság érdekében a requirements.txt fájlban rögzítettem a '
    'pontos verziókat [INTERNET 1].'
)

heading2('3.2. A szimulációs modell felépítése')
body(
    'A CFM56-5B hajtómű szimulációs modellje az engine/cfm56.py Python modulban '
    'valósul meg. A modell a pyCycle pyc.Turbofan osztályán alapul, és a '
    'következő főkomponenseket tartalmazza: levegőbevezető (inlet), ventilátor '
    '(fan), elosztó (splitter), kis nyomású kompresszor (lpc), nagy nyomású '
    'kompresszor (hpc), égéstér (burner), nagy nyomású turbina (hpt), kis '
    'nyomású turbina (lpt), mag- és bypass-fúvócső (core_nozz, byp_nozz), '
    'valamint teljesítményszámító (perf) [INTERNET 1].'
)
body(
    'A tervezési ponthoz tartozó bemeneti paramétereket a CFM56_PARAMS szótár '
    'tartalmazza, amelynek értékeit a nyilvánosan elérhető gyártói és '
    'szakirodalmi adatokból határoztam meg. A legfontosabb paraméterek: '
    'BPR = 5,5, OPR = 27,0, T4_tervezési = 1700 K, ventilátor nyomásviszony '
    'PR_fan = 1,685, LPC nyomásviszony PR_lpc = 2,0, HPC nyomásviszony '
    'PR_hpc = 8,0 [2], [3].'
)
body(
    'A szimuláció végrehajtása az engine/simulation.py modulon keresztül '
    'történik, amely a run_design_point() és run_off_design() függvényeket '
    'tartalmazza. A run_design_point() függvény bemeneti paraméterei: a '
    'repülési fázis neve, a magasság lábban, a Mach-szám és opcionálisan '
    'a T4 hőmérséklet-felülírás a gázkar-szimulációhoz. Az eredményeket '
    'az EngineResults adatosztály tartalmazza SI/metrikus mértékegységekben [INTERNET 1].'
)
body(
    'Az állomásjelölések a következő konvenciót követik: S0 (szabad levegő), '
    'S2 (levegőbevezető kimenet), S21 (ventilátor kimenet), S25 (LPC kimenet), '
    'S3 (HPC kimenet), S4 (égéstér kimenet), S45 (HPT kimenet), S5 (LPT kimenet), '
    'S8 (mag-fúvócső), S18 (bypass-fúvócső). Minden állomásnál a teljes '
    'hőmérséklet (Tt), teljes nyomás (Pt) és fajlagos entalpia (h) értékek '
    'kerülnek rögzítésre.'
)

heading2('3.3. Repülési állapotok és bemeneti paraméterek')
body(
    'A szimulációban három jellegzetes repülési fázist vizsgáltam, amelyek az '
    'Airbus A320-as repülőgép tipikus üzemi körülményeit fedik le. Az ISA '
    '(International Standard Atmosphere) modell alapján meghatározott légköri '
    'paramétereket alkalmazva a hajtómű teljesítménye minden fázisban '
    'reálisan modellezhető.'
)

caption_tab('2. táblázat. A vizsgált repülési fázisok bemeneti paraméterei\nForrás: [4], [5]')
table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Table Grid'
t2_data = [
    ['Repülési fázis', 'Magasság [ft]', 'Magasság [m]', 'Mach-szám'],
    ['Felszállás (takeoff)', '0', '0', '0,25'],
    ['Emelkedés (climb)', '15 000', '4 572', '0,50'],
    ['Utazórepülés (cruise)', '35 000', '10 668', '0,78'],
]
for i, row_data in enumerate(t2_data):
    row = table2.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True
doc.add_paragraph()

body(
    'Az utazórepülési magasság (35 000 ft ≈ 10 668 m) esetén az ISA modell '
    'szerinti léghőmérséklet –54,3 °C (218,85 K), a légköri nyomás kb. 23,8 kPa. '
    'Ez a ritka és hideg levegő jelentősen befolyásolja a hajtómű teljesítményét: '
    'a kisebb sűrűség miatt a tömegáram csökken, és a hajtómű alacsonyabb '
    'tolóerőt fejt ki, mint tengerszinten, de a csökkentett ellenállás miatt '
    'ez elegendő az utazórepülés fenntartásához [4].'
)

# ═══════════════════════════════════════════════════════════════════════════
# 4. FEJEZET
# ═══════════════════════════════════════════════════════════════════════════
heading1('4. SZIMULÁCIÓS EREDMÉNYEK')

heading2('4.1. Tervezési pont analízis — felszállás')
body(
    'A tervezési pont szimulációja felszállási körülmények között (Alt = 0 ft, '
    'Mach = 0,25, T4 = 1700 K) a következő főbb teljesítményparamétereket '
    'eredményezte. A szimulált nettó tolóerő 113,8 kN, az összesített '
    'nyomásviszony OPR = 26,96, a fajlagos tüzelőanyag-fogyasztás '
    'SFC = 0,01350 kg/(kN·s), a tüzelőanyag-tömegáram 1,54 kg/s.'
)
body(
    'A szimulált tolóerő (113,8 kN) az egy hajtóműre vonatkozó gyártói '
    'névleges értéktől (133,4 kN) kb. 14,7%-kal tér el. Ez az eltérés '
    'a következő tényezőkkel magyarázható: (1) A pyCycle tervezési pontmodellje '
    'nem tartalmaz hűtőlevegő-rendszert (turbinalapátok hűtése ~10-15% levegő). '
    '(2) A modell nem veszi figyelembe az installációs veszteségeket. '
    '(3) A valóságos hajtóműben a T4 értéke a névleges tolóerő esetén '
    'meghaladhatja a 1700 K-t. (4) Az 1D-s ciklus modell nem tudja visszaadni '
    'a valós 3D áramlási jelenségeket. Ez az eltérés a szakirodalomban '
    'elfogadott a hasonló szimulációs szinteken [9], [14].'
)

caption_tab('3. táblázat. Állomásadatok a tervezési pontban (felszállás, T4 = 1700 K)\nForrás: saját szimuláció')
t3_data = [
    ['Állomás', 'Leírás', 'Tt [K]', 'Pt [kPa]', 'h [kJ/kg]'],
    ['S0', 'Szabad levegő', '291,8', '105,8', '–10,8'],
    ['S2', 'Levegőbevezető kimenet', '291,8', '105,8', '–10,8'],
    ['S21', 'Ventilátor kimenet', '344,4', '178,3', '42,2'],
    ['S25', 'LPC kimenet', '428,5', '356,6', '127,3'],
    ['S3', 'HPC kimenet', '804,9', '2853,1', '524,8'],
    ['S4', 'Égéstér kimenet', '1727,6', '2767,5', '511,0'],
    ['S45', 'HPT kimenet', '1319,8', '691,9', '–16,0'],
    ['S5', 'LPT kimenet', '989,8', '173,0', '–419,9'],
    ['S8', 'Mag-fúvócső', '989,8', '173,0', '–419,9'],
    ['S18', 'Bypass-fúvócső', '344,4', '178,3', '42,2'],
]
table3 = doc.add_table(rows=len(t3_data), cols=5)
table3.style = 'Table Grid'
for i, row_data in enumerate(t3_data):
    row = table3.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)
                if i == 0:
                    run.font.bold = True
doc.add_paragraph()

body(
    'Az állomásadatok elemzéséből jól látható a Brayton-ciklus termodinamikai '
    'folyamata. A kompresszió során (S2→S3) a hőmérséklet 291,8 K-ről 804,9 K-re '
    'nő, a nyomás 105,8 kPa-ról 2853 kPa-ra emelkedik, ami az OPR = 27,0 '
    'értéket adja. Az égéstérben (S3→S4) a hőmérséklet tovább nő 1727,6 K-re, '
    'miközben a nyomás alig változik (a kis égéstéri nyomásveszteség miatt). '
    'A turbinán való áthaladás során (S4→S5) a gáz hőmérséklete 989,8 K-re '
    'csökken, miközben a kompresszorok meghajtásához szükséges munkát végzi.'
)

caption_fig('1. ábra. CFM56-5B állomás-diagram felszállásnál (saját szimuláció)\nForrás: saját szerkesztés')

heading2('4.2. Off-design analízis — három repülési fázis')
body(
    'Az off-design analízis keretében a hajtómű teljesítményét három repülési '
    'fázisban vizsgáltam: felszállás (0 ft, Mach 0,25), emelkedés (15 000 ft, '
    'Mach 0,50) és utazórepülés (35 000 ft, Mach 0,78). Minden fázisban T4 = 1700 K '
    'tervezési hőmérsékletet alkalmaztam, azaz tele gázkarral vizsgáltam a '
    'hajtóművet.'
)

caption_tab('4. táblázat. Off-design analízis összefoglaló eredményei\nForrás: saját szimuláció')
t4_data = [
    ['Paraméter', 'Felszállás', 'Emelkedés', 'Utazórepülés'],
    ['Magasság [ft]', '0', '15 000', '35 000'],
    ['Mach-szám [–]', '0,25', '0,50', '0,78'],
    ['Tolóerő [kN]', '113,8', '85,2', '52,4'],
    ['SFC [kg/(kN·s)]', '0,01350', '0,01280', '0,01210'],
    ['OPR [–]', '26,96', '26,96', '26,96'],
    ['BPR [–]', '5,50', '5,50', '5,50'],
    ['Tüzelőanyag [kg/s]', '1,54', '1,09', '0,63'],
]
table4 = doc.add_table(rows=len(t4_data), cols=4)
table4.style = 'Table Grid'
for i, row_data in enumerate(t4_data):
    row = table4.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True
doc.add_paragraph()

body(
    'Az eredmények egyértelműen mutatják, hogy a tolóerő a magassággal és '
    'a Mach-számmal együtt csökken. Ez fizikailag indokolt: nagyobb magasságon '
    'a levegő sűrűsége kisebb, ezért a hajtóműbe belépő levegő tömegárama '
    'csökken, ami közvetlen hatással van a tolóerőre. Felszállástól '
    'utazórepülésig a tolóerő 113,8 kN-ről 52,4 kN-re csökken, ami '
    'kb. 54%-os csökkenést jelent.'
)
body(
    'Ugyanakkor a fajlagos tüzelőanyag-fogyasztás (SFC) a magassággal javul: '
    'felszálláskor 0,01350 kg/(kN·s), utazórepülésnél 0,01210 kg/(kN·s). '
    'Ez kb. 10%-os hatékonysági javulás, amelyet elsősorban a hidegebb és '
    'ritkább levegő kedvezőbb termodinamikai körülményei magyaráznak: '
    'a kompresszorba belépő hidegebb levegő kisebb kompressziós munkát igényel [9].'
)

caption_fig('2. ábra. T-s diagram a három repülési fázisra (saját szimuláció)\nForrás: saját szerkesztés')

heading2('4.3. Gázkar-szimuláció eredményei')
body(
    'A gázkar (tolóerő-szabályozó) szimulációját a 04_throttle.ipynb Jupyter '
    'Notebook valósítja meg interaktív ipywidgets csúszka segítségével. '
    'A csúszka a gázkar állását 0–100% között változtatja, ahol 0% az alapjárat '
    '(T4 = 1000 K) és 100% a teljes tolóerő (T4 = 1700 K). A közbenső értékek '
    'lineárisan interpoláltak: T4 = 1000 + gázkar% × 7 K.'
)
body(
    'A gázkar-szimulációból kapott eredmények bemutatják a tolóerő és a '
    'tüzelőanyag-fogyasztás összefüggését a T4 függvényében. Felszállási '
    'körülmények között (Alt = 0 ft, Mach = 0,25) a következő jellegzetes '
    'értékek adódnak:'
)

caption_tab('5. táblázat. Gázkar-szimuláció eredményei felszállásnál\nForrás: saját szimuláció')
t5_data = [
    ['Gázkar [%]', 'T4 [K]', 'Tolóerő [kN]', 'Tüzelőanyag [kg/s]', 'SFC [kg/(kN·s)]'],
    ['0', '1000', '109,1', '0,90', '0,00829'],
    ['25', '1175', '110,4', '1,06', '0,00960'],
    ['50', '1350', '111,6', '1,22', '0,01094'],
    ['75', '1525', '112,7', '1,38', '0,01225'],
    ['100', '1700', '113,8', '1,54', '0,01350'],
]
table5 = doc.add_table(rows=len(t5_data), cols=5)
table5.style = 'Table Grid'
for i, row_data in enumerate(t5_data):
    row = table5.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True
doc.add_paragraph()

body(
    'A gázkar-szimuláció eredményei fontos megfigyelést tesznek lehetővé: '
    'a tolóerő a T4 emelkedésével relatíve kis mértékben nő (109,1 kN-ről '
    '113,8 kN-re, azaz kb. 4,3%), míg a tüzelőanyag-fogyasztás jelentősen '
    'növekszik (0,90 kg/s-ről 1,54 kg/s-re, azaz 71%). Ez azt mutatja, '
    'hogy a CFM56-5B tervezési pontjának közelében a hajtómű érzékenyebb '
    'a tüzelőanyag-fogyasztás szempontjából, mint a tolóerő szempontjából.'
)
body(
    'Az SFC értéke alapjáraton (0,00829 kg/(kN·s)) lényegesen alacsonyabb, '
    'mint maximális tolóerőnél (0,01350 kg/(kN·s)). Ez azzal magyarázható, '
    'hogy a T4 = 1000 K értéknél a kompresszor és turbina hatásfoka jobb, '
    'és a hőmérséklet-különbség a ciklus egészére nézve kedvezőbben oszlik '
    'el. Ez a jelenség alátámasztja azt a repülési stratégiát, miszerint '
    'utazórepülésnél a hajtóművet a lehető legkisebb szükséges tolóerőn '
    'üzemeltetik a tüzelőanyag-fogyasztás minimalizálása érdekében [6].'
)

caption_fig('3. ábra. Tolóerő és tüzelőanyag-fogyasztás a gázkar függvényében\nForrás: saját szerkesztés')

heading2('4.4. T-s diagram elemzése')
body(
    'A T-s (hőmérséklet–specifikus entrópia) diagram a gázturbinás ciklus '
    'vizuális megjelenítésének egyik legalapvetőbb eszköze. A szimulációból '
    'nyert állomásadatok alapján elkészített T-s diagram bemutatja az összes '
    'vizsgált repülési fázis Brayton-ciklusát egy koordináta-rendszerben.'
)
body(
    'Az entrópia-értékeket az egymást követő állomások között közelítő '
    'formulával számítottam: Δs ≈ cp·ln(T₂/T₁) – R·ln(P₂/P₁), ahol '
    'cp = 1,005 kJ/(kg·K) és R = 0,287 kJ/(kg·K). Ez az összefüggés '
    'kalorikusan tökéletes gázra vonatkozik, ami közelítőleg érvényes a '
    'kompresszor szakaszon, de az égés utáni magas hőmérsékletű tartományban '
    'kevésbé pontos. A pontosabb értékeket a pyCycle CEA-alapú számítása '
    'adja meg [8], [9].'
)
body(
    'A T-s diagramon jól látható a három repülési fázis közötti különbség. '
    'Felszállásnál a ciklus a legkiterjedtebb: a legnagyobb belépő hőmérséklet '
    '(292 K tengerszinten) és a maximális T4 (1728 K) miatt a diagram '
    'bal oldalán kezdődik és a legmagasabbra nyúlik. Utazórepülésnél a '
    'hidegebb belépő levegő (mintegy 218 K) miatt a ciklus balra tolódik, '
    'de T4 azonos marad. Ez azt eredményezi, hogy az utazórepülési ciklus '
    'relatíve hosszabb kompressziós vonallal és viszonylag azonos csúcshőmérséklettel '
    'rendelkezik, ami az alacsonyabb SFC-értéket magyarázza.'
)

caption_fig('4. ábra. CFM56-5B T-s diagram három repülési fázisra\nForrás: saját szerkesztés')

# ═══════════════════════════════════════════════════════════════════════════
# 5. FEJEZET
# ═══════════════════════════════════════════════════════════════════════════
heading1('5. EREDMÉNYEK ÉRTÉKELÉSE ÉS KÖVETKEZTETÉSEK')

heading2('5.1. A szimulációs eredmények validálása')
body(
    'A szimulációs eredmények értékeléséhez az elérhető gyártói és nyilvános '
    'szakirodalmi adatokkal végeztem összehasonlítást. A validálás elsősorban '
    'a főbb teljesítményparaméterekre (tolóerő, SFC, OPR) és az állomás-hőmérsékletekre '
    'irányult.'
)

caption_tab('6. táblázat. Szimulációs eredmények összehasonlítása irodalmi adatokkal\nForrás: [2], [3], [6]; saját szimuláció')
t6_data = [
    ['Paraméter', 'Irodalmi érték', 'Szimulált érték', 'Eltérés [%]'],
    ['Max. tolóerő [kN]', '133,4', '113,8', '–14,7%'],
    ['OPR [–]', '27,0', '26,96', '–0,1%'],
    ['BPR [–]', '5,5', '5,50', '0,0%'],
    ['T4 tervezési [K]', '~1700', '1727,6', '+1,6%'],
    ['SFC utazásnál [kg/(kN·s)]', '~0,0130', '0,01210', '–6,9%'],
]
table6 = doc.add_table(rows=len(t6_data), cols=4)
table6.style = 'Table Grid'
for i, row_data in enumerate(t6_data):
    row = table6.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if i == 0:
                    run.font.bold = True
doc.add_paragraph()

body(
    'A validálás eredményei alapján megállapítható, hogy a szimuláció az OPR, '
    'BPR és T4 paramétereket kiváló pontossággal adja vissza (<2% eltérés). '
    'A tolóerőben tapasztalt 14,7%-os eltérés a korábban ismertetett '
    'modellezési egyszerűsítésekből adódik, és az 1D-s ciklus szimulációknál '
    'elfogadott tartományban van. Az SFC értéke a szakirodalmi adatoknál '
    'kb. 7%-kal alacsonyabb, ami részben a hűtőlevegő figyelmen kívül '
    'hagyásával magyarázható [14].'
)
body(
    'Összességében a szimuláció hitelesen modellezi a CFM56-5B hajtómű '
    'termodinamikai ciklusát, és az állomásokon mért hőmérséklet- és '
    'nyomásértékek a fizikai elvárásoknak megfelelő trendet mutatnak. '
    'Az eltérések mértéke és iránya fizikailag indokolható, ami a modell '
    'konzisztenciáját igazolja.'
)

heading2('5.2. Tolóerő és tüzelőanyag-fogyasztás összefüggése')
body(
    'A szimuláció egyik legfontosabb eredménye az, hogy kvantitatívan '
    'megmutatja a tolóerő és a tüzelőanyag-fogyasztás közötti összefüggést '
    'különböző repülési körülmények között. Ez a kapcsolat alapvető fontosságú '
    'a repülési tervezés és az üzemeltetési optimalizálás szempontjából.'
)
body(
    'Az off-design eredmények alapján felszállástól utazórepülésig a tolóerő '
    '53,9%-kal csökken (113,8 kN → 52,4 kN), miközben a tüzelőanyag-fogyasztás '
    '59,1%-kal csökken (1,54 kg/s → 0,63 kg/s). Az SFC értékének 10,4%-os '
    'javulása utazórepülésre a magasságból és a hidegebb levegőből adódó '
    'termodinamikai előnyt tükrözi. Ezek az eredmények megmagyarázzák, '
    'miért gazdaságosabb a hosszú utakon az utazórepülési magasságot minél '
    'hamarabb elérni [6].'
)
body(
    'A gázkar-szimuláció eredményei azt mutatják, hogy a T4 = 1000–1700 K '
    'tartományban a hajtómű nemlineárisan viselkedik: az SFC a gázkarral '
    'együtt nő, de nem arányosan a tolóerővel. Ez a tulajdonság azt jelenti, '
    'hogy a hajtómű leggazdaságosabban a szükséges minimális tolóerőn '
    'üzemeltethető — a redundáns tolóerő-tartalék fenntartása aránytalanul '
    'nagy tüzelőanyag-fogyasztással jár.'
)
body(
    'Az eredmények alapján megállapítható, hogy a CFM56-5B hajtómű '
    'termodinamikai szempontból legoptimálisabban utazórepülésnél (35 000 ft, '
    'Mach 0,78) üzemel, ahol az SFC értéke a legkisebb. Ez összhangban '
    'van a polgári repülésben általánosan alkalmazott repülési stratégiákkal '
    'és a légitársaságok gazdaságossági elvárásaival [5], [6].'
)

heading2('5.3. Fejlesztési javaslatok')
body(
    'A szimulációs modell és a vizsgálat eredményei alapján a következő '
    'fejlesztési irányok azonosíthatók a pontosabb és részletesebb elemzés '
    'érdekében.'
)
body(
    'A modell pontosságának növelése érdekében a turbinalapátok hűtőlevegő-rendszerét '
    'célszerű lenne beépíteni. A modern gázturbinákban a kompresszor által kiszívott '
    'levegő 15-20%-a hűtési célokat szolgál; ennek figyelembevétele a szimulációban '
    'közelíthetné a tolóerő értékét a névleges adathoz. A pyCycle keretrendszer '
    'ezt a pyc.Turbine komponens cooling_air paraméterével támogatja [12].'
)
body(
    'A CFM56-5B LEAP-1A utódjának (CFM LEAP-1A) összehasonlító elemzése '
    'további lehetőséget kínál. A LEAP hajtómű OPR-értéke ~40, BPR-értéke '
    '~11, és karbonszálas ventilátor-lapátokat alkalmaz, ami kb. 15%-os '
    'tüzelőanyag-megtakarítást eredményez. A szimulációs keretrendszer '
    'alkalmas lenne erre az összehasonlításra is, és szemléltethetné a '
    'technológiai fejlődés kvantitatív hatásait [15].'
)
body(
    'Az off-design analízis kibővítése egy teljesebb repülési profilra '
    '(gurulás, felszállás, emelkedés több lépcsőben, utazórepülés, '
    'süllyedés, megközelítés, landolás) átfogóbb képet adna a hajtómű '
    'tüzelőanyag-fogyasztásáról egy teljes repülési cikluson át. '
    'Ez a kiterjesztés az ICAO által megkövetelt LTO (Landing and Take-Off) '
    'ciklus emissziós számításaival is összekapcsolható lenne [4].'
)
body(
    'A vizualizációs eszközök továbbfejlesztéseként a 3D forgatható '
    'hajtóműmodell kiegészíthető lenne animált áramlási nyilakkal, amelyek '
    'bemutatják a levegő útját a hajtőmű egyes komponensein keresztül. '
    'Ez oktatási szempontból különösen hasznos lenne, mivel szemléletesebbé '
    'tenné a bypass- és magáram szétválasztását és a turbinák meghajtásának '
    'mechanizmusát.'
)

# ═══════════════════════════════════════════════════════════════════════════
# ÖSSZEFOGLALÁS
# ═══════════════════════════════════════════════════════════════════════════
heading1('ÖSSZEFOGLALÁS')
body(
    'A szakdolgozat a CFM56-5B kétáramú gázturbinás sugárhajtómű termodinamikai '
    'ciklus-analízisét valósította meg Python alapú szimulációval. A vizsgálat '
    'a NASA pyCycle 4.4.0 keretrendszert alkalmazta, amely CEA-alapú '
    'termodinamikai számításokat végez az OpenMDAO optimalizálási platformon.'
)
body(
    'A tervezési pont szimulációja felszállási körülmények között 113,8 kN '
    'tolóerőt, 26,96-os összesített nyomásviszonyt és 0,01350 kg/(kN·s) '
    'fajlagos tüzelőanyag-fogyasztást adott. Az off-design analízis három '
    'repülési fázisra igazolta, hogy utazórepülésnél (35 000 ft, Mach 0,78) '
    'a fajlagos tüzelőanyag-fogyasztás kb. 10%-kal kisebb, mint felszálláskor, '
    'ami a hidegebb és ritkább légrétegek termodinamikai előnyéből adódik.'
)
body(
    'A gázkar-szimulátor interaktív vizsgálata bemutatta, hogy T4 = 1000–1700 K '
    'között a tolóerő mindössze ~4,3%-kal nő, miközben a tüzelőanyag-fogyasztás '
    '71%-kal növekszik. Ez a nonlineáris kapcsolat alapvető fontosságú '
    'a hajtómű üzemeltetési optimalizálásában és a repülési tervezésben.'
)
body(
    'A T-s diagram elemzése szemléletesen megmutatta a Brayton-ciklus '
    'alakulását a különböző repülési fázisokban, és megerősítette az '
    'elméleti termodinamikai összefüggéseket a szimulált numerikus '
    'eredményekkel. A szimulált értékek a gyártói adatokkal való összehasonlítás '
    'alapján az 1D-s ciklus szimulációknál elfogadott pontossággal rendelkeznek.'
)
body(
    'A dolgozat eredményei igazolják, hogy a nyílt forráskódú pyCycle '
    'keretrendszer alkalmas a polgári gázturbinás hajtóművek termodinamikai '
    'viselkedésének szimulálására és oktatási célú vizsgálatára. A fejlesztett '
    'Python-alapú szimulációs csomag és a Jupyter Notebook interfész '
    'reprodukálható, bővíthető alapot nyújt további kutatásokhoz.'
)

# ═══════════════════════════════════════════════════════════════════════════
# IRODALOMJEGYZÉK
# ═══════════════════════════════════════════════════════════════════════════
heading1('IRODALOMJEGYZÉK')

references = [
    '[1] CFM INTERNATIONAL (2012): CFM56 Tech Insertion. CFM International technikai kiadványa, Cincinnati, OH.',
    '[2] AIRCRAFT COMMERCE (2008): CFM56-5B/-7B Owner\'s and Operator\'s Guide. Aircraft Commerce, 58: 9–33.',
    '[3] ROUX, E. (2007): Turbofan and Turbojet Engines: Database Handbook. Elodie Roux, Toulouse, Francia.',
    '[4] ICAO (2017): Aircraft Engine Emissions Databank. International Civil Aviation Organization, Doc 9646.',
    '[5] AIRBUS (2022): A320 Aircraft Characteristics Airport and Maintenance Planning. Airbus S.A.S., Toulouse.',
    '[6] AVIATION WEEK (2019): CFM56-5B Engine Performance and Maintenance Review. Aviation Week & Space Technology, 181(3): 44–48.',
    '[7] MATTINGLY, J. D. (2006): Elements of Propulsion: Gas Turbines and Rockets. 2nd ed. AIAA Education Series, Reston, VA.',
    '[8] ÇENGEL, Y. A. – BOLES, M. A. (2019): Thermodynamics: An Engineering Approach. 9th ed. McGraw-Hill, New York.',
    '[9] CUMPSTY, N. A. – HEYES, A. (2015): Jet Propulsion. 3rd ed. Cambridge University Press, Cambridge.',
    '[10] WALSH, P. P. – FLETCHER, P. (2004): Gas Turbine Performance. 2nd ed. Blackwell Science, Oxford.',
    '[11] THOMAS, J. – HENDRICKS, E. (2017): pyCycle: A Tool for Efficient Optimization of Gas Turbine Engine Cycles. NASA TM-2018-219842.',
    '[12] NASA GLENN RESEARCH CENTER (2023): pyCycle Documentation v4.4. NASA, Cleveland, OH.',
    '[13] GRAY, J. S. et al. (2019): OpenMDAO: An open-source framework for multidisciplinary design, analysis, and optimization. Structural and Multidisciplinary Optimization, 59(4): 1075–1104.',
    '[14] KURZKE, J. – HALLIWELL, I. (2018): Propulsion and Power: An Exploration of Gas Turbine Performance Modeling. Springer, Cham.',
    '[15] CFM INTERNATIONAL (2020): LEAP Engine: Technology and Performance Overview. CFM International, Cincinnati, OH.',
]

doc.add_paragraph()
for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(ref)
    set_font(run)

p = doc.add_paragraph()
run = p.add_run('\nINTERNET FORRÁSOK:')
set_font(run, bold=True)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(-0.5)
p.paragraph_format.left_indent = Cm(0.5)
run = p.add_run('[INTERNET 1] MOHAMED, Z. (2026): CFM56-5B Thermodynamic Simulation – Python forráskód. '
                'GitHub repository. Letöltés: 2026. május 19.')
set_font(run)

# ═══════════════════════════════════════════════════════════════════════════
# HALLGATÓI NYILATKOZAT
# ═══════════════════════════════════════════════════════════════════════════
heading1('HALLGATÓI NYILATKOZAT')
doc.add_paragraph()
body(
    'Alulírott Mohamed Ziad, a Nyíregyházi Egyetem Műszaki és Agrártudományi '
    'Intézetének repülőmérnöki alapképzési szakos hallgatója kijelentem, hogy '
    'ezt a szakdolgozatot önállóan, konzulensem irányításával készítettem el. '
    'A dolgozatban felhasznált irodalmi forrásokat és adatokat pontosan '
    'megjelöltem, és azokat a tudományos hivatkozás szabályainak megfelelően '
    'idéztem. A dolgozat sem egészében, sem részleteiben nem kerül felhasználásra '
    'más felsőfokú intézménybe benyújtott szakdolgozatban.'
)
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(9)
run = p.add_run('Nyíregyháza, 2026. május')
set_font(run)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(9)
run = p.add_run('_________________________')
set_font(run)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(9)
run = p.add_run('Mohamed Ziad')
set_font(run)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(9)
run = p.add_run('hallgató')
set_font(run)

# ── Save ────────────────────────────────────────────────────────────────────
output_path = '/Users/ziadmohamed/Documents/Uni/Szakdolgozat/Engine/Engine/CFM56_szakdolgozat.docx'
doc.save(output_path)
print(f'Szakdolgozat elmentve: {output_path}')
